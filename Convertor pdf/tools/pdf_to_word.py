import os
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTTextLine
from PyPDF2 import PdfReader

def handle_pdf_to_word(file):
    """
    Advanced: Convert an uploaded PDF file to a Word (.docx) document, preserving basic formatting (bold, italic, font size, color).
    :param file: FileStorage object (uploaded PDF)
    :return: path to generated .docx file
    """
    # Remove debug import and print for production

    # Directories
    UPLOAD_DIR = 'uploads'
    OUTPUT_DIR = 'output'
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Save PDF to disk
    original_name = file.filename
    if not original_name.lower().endswith('.pdf'):
        raise Exception("Unsupported file type. Please upload a PDF file.")
    input_path = os.path.join(UPLOAD_DIR, original_name)
    file.save(input_path)
    # Debug: print file size after save
    print(f"DEBUG: Saved file size: {os.path.getsize(input_path)} bytes")
    # Check if file is empty
    if os.path.getsize(input_path) == 0:
        os.remove(input_path)
        raise Exception("The uploaded file is empty. Please select a valid PDF file.")

    # Validate PDF before processing
    try:
        with open(input_path, 'rb') as f:
            PdfReader(f)
    except Exception:
        os.remove(input_path)
        raise Exception("The file you uploaded is not a valid PDF. Please upload a real PDF file.")

    # Create a Word document
    doc = Document()

    # Open PDF and extract text
    try:
        for page_layout in extract_pages(input_path):
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    for text_line in element:
                        if isinstance(text_line, LTTextLine):
                            para = doc.add_paragraph()
                            last_run = None
                            for char in text_line:
                                if isinstance(char, LTChar):
                                    run = para.add_run(char.get_text())
                                    font = run.font
                                    # Font size
                                    if char.size:
                                        font.size = Pt(char.size)
                                    # Bold/Italic
                                    font.bold = 'Bold' in char.fontname
                                    font.italic = 'Italic' in char.fontname or 'Oblique' in char.fontname
                                    # Color (if available)
                                    if hasattr(char, 'graphicstate') and hasattr(char.graphicstate, 'stroking_color'):
                                        color = char.graphicstate.stroking_color
                                        if isinstance(color, tuple) and len(color) == 3:
                                            r, g, b = [int(255 * c) for c in color]
                                            font.color.rgb = RGBColor(r, g, b)
                                    last_run = run
                                else:
                                    # For LTAnno (spaces, etc.)
                                    if last_run:
                                        last_run.add_text(char.get_text())
                            para.alignment = 0  # Left align
            doc.add_page_break()
    except Exception as e:
        os.remove(input_path)
        raise Exception(f"Failed to extract text and formatting from PDF: {str(e)}")

    # Save .docx
    output_filename = f"{os.path.splitext(original_name)[0]}_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)

    # Cleanup upload
    os.remove(input_path)

    return output_path
